import docx
import sys

doc = docx.Document(r'try on.docx')

with open('docx_content.txt', 'w', encoding='utf-8') as f:
    f.write("=== PARAGRAPHS ===\n")
    for i, p in enumerate(doc.paragraphs):
        f.write(f"[{i}] {repr(p.text)}\n")
    f.write("\n=== TABLES ===\n")
    for ti, table in enumerate(doc.tables):
        f.write(f"--- Table {ti} ---\n")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            f.write(" | ".join(cells) + "\n")

print("Done - wrote to docx_content.txt")
